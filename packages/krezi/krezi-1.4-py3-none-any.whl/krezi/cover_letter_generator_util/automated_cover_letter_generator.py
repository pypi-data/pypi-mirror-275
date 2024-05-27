from docx import Document
from datetime import datetime
from docx import Document
import argparse

from docx.enum.text import *
from docx.shared import Pt, Cm

date_suffix_dict = {}
date_suffix_dict.update({i:'st' for i in [1, 21, 31]})
date_suffix_dict.update({i:'nd' for i in [2, 12, 22]})

parser = argparse.ArgumentParser(
                    prog='Cover Letter Generator',
                    description='With a few variables, generates a pre worded cover letter')

parser.add_argument("-c", "--company", required = True,
                    help = "company you are applying for")
parser.add_argument("-r", "--role-name", required = True, 
                    help = "role for which you are applying for")
parser.add_argument("-s", "--source",  required = False,
                    default = 'LinkedIn', 
                    help = "where you found this role : Ex LinkedIn")
parser.add_argument("-i", "--is-internship", 
                    action='store_true', 
                    help = "is the role an internship")

args = parser.parse_args()

DATE_TODAY = datetime.now()

ROLE_NAME = args.role_name # role r

IS_INTERNSHIP = args.is_internship # internship i

COMPANY_NAME = args.company # company c

SOURCE = args.source # source s

MARGIN = 1.27 # Taken from Word : Narrow Margin 
LINE_SPACING = 1.15 # Taken from Word

if __name__ == '__main__':
    d = Document()

    header_line_1 = "Your Name"
    header_line_2 = "location | Phone | LinkedIn Url | Github Url"

    header = d.sections[0].header

    head_p0 = header.paragraphs[0]
    head_p0.text = header_line_1
    pf = head_p0.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head_p0.runs[0].bold = True
    fhp0 = head_p0.runs[0].font
    fhp0.size = Pt(18)
    fhp0.name = "Times New Roman"

    head_p1 = header.add_paragraph()
    head_p1.text = header_line_2
    pf = head_p1.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head_p1.runs[0].bold = False
    fhp1 = head_p1.runs[0].font
    fhp1.size = Pt(9)
    fhp1.name = "Times New Roman"

    paragraph_1 = "Subject"
    paragraph_2 = "Dear Manager,"
    paragraph_3 = "p3_text"
    paragraph_4 = "p4_text"
    paragraph_5 = "p5_text"
    paragraph_6 = "p6_text"
    paragraph_7 = "p6_text"
    paragraph_8 = "p7_text"
    ending_message = "Sincerely,\nYour Name"

    paragraph_1 = paragraph_1.replace(":role_name", ROLE_NAME)
    paragraph_3 = paragraph_3.replace(":role_name", ROLE_NAME)
    paragraph_3 = paragraph_3.replace(":source", SOURCE if SOURCE else "LinkedIn")
    paragraph_3 = paragraph_3.replace(":company_name", COMPANY_NAME)
    paragraph_7 = paragraph_7.replace(":company_name", COMPANY_NAME)

    p1 = d.add_paragraph(paragraph_1)
    fbp1 = p1.paragraph_format
    fbp1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fbp1.space_before = Pt(10)
    fbp1.space_after = Pt(10)
    fbp1.line_spacing = 1.15
    fbp1 = p1.runs[0].font
    fbp1.size = Pt(10)
    fbp1.name = "Times New Roman"

    p2 = d.add_paragraph(paragraph_2)
    fbp2 = p2.paragraph_format
    fbp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fbp2.space_after = Pt(10)
    fbp2.line_spacing = 1.15
    fbp2 = p2.runs[0].font
    fbp2.size = Pt(10)
    fbp2.name = "Times New Roman"

    p3 = d.add_paragraph(paragraph_3)
    fbp3 = p3.paragraph_format
    fbp3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fbp3.space_after = Pt(10)
    fbp3.line_spacing = 1.15
    fbp3 = p3.runs[0].font
    fbp3.size = Pt(10)
    fbp3.name = "Times New Roman"

    p4 = d.add_paragraph(paragraph_4)
    fbp4 = p4.paragraph_format
    fbp4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fbp4.space_after = Pt(10)
    fbp4.line_spacing = 1.15
    fbp4 = p4.runs[0].font
    fbp4.size = Pt(10)
    fbp4.name = "Times New Roman"

    p5 = d.add_paragraph(paragraph_5)
    fbp5 = p5.paragraph_format
    fbp5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fbp5.space_after = Pt(10)
    fbp5.line_spacing = 1.15
    fbp5 = p5.runs[0].font
    fbp5.size = Pt(10)
    fbp5.name = "Times New Roman"

    p6 = d.add_paragraph(paragraph_6)
    fbp6 = p6.paragraph_format
    fbp6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fbp6.space_after = Pt(10)
    fbp6.line_spacing = 1.15
    fbp6 = p6.runs[0].font
    fbp6.size = Pt(10)
    fbp6.name = "Times New Roman"

    p7 = d.add_paragraph(paragraph_7)
    fbp7 = p7.paragraph_format
    fbp7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fbp7.space_after = Pt(10)
    fbp7.line_spacing = 1.15
    fbp7 = p7.runs[0].font
    fbp7.size = Pt(10)
    fbp7.name = "Times New Roman"

    p8 = d.add_paragraph(paragraph_8)
    fbp8 = p8.paragraph_format
    fbp8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fbp8.space_after = Pt(10)
    fbp8.line_spacing = 1.15
    fbp8 = p8.runs[0].font
    fbp8.size = Pt(10)
    fbp8.name = "Times New Roman"

    _ending_message = d.add_paragraph(ending_message)
    fb_ending_message = _ending_message.paragraph_format
    fb_ending_message.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fb_ending_message.space_after = Pt(10)
    fb_ending_message.line_spacing = 1.15
    fb_ending_message = _ending_message.runs[0].font
    fb_ending_message.size = Pt(10)
    fb_ending_message.name = "Times New Roman"


    sections = d.sections
    for section in sections:
        section.top_margin = Cm(MARGIN)
        section.bottom_margin = Cm(MARGIN)
        section.left_margin = Cm(MARGIN)
        section.right_margin = Cm(MARGIN)

    d.save(f"/location_for_saving_file/Cover Letter {COMPANY_NAME} {ROLE_NAME} {DATE_TODAY:%Y-%m-%d}.docx")
